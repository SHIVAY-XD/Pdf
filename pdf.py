from docx import Document

# Create a new Document
doc = Document()

# Add a title to the document
doc.add_heading('Python Programming - Question Bank with Answers', 0)

# Unit 1 Questions and Answers
doc.add_heading('Unit 1', level=1)

# Question 1
doc.add_heading('1. What is Python? List and explain features of Python.', level=2)
doc.add_paragraph(
    "Python is a high-level, interpreted, and general-purpose programming language. "
    "It was created by Guido van Rossum and first released in 1991. Python is known for its simplicity, readability, and ease of use. "
    "It supports multiple programming paradigms, such as procedural, object-oriented, and functional programming.\n\n"
    "Features of Python include:\n"
    "1. **Easy to learn and use**: Python has a simple syntax, which makes it easy for beginners.\n"
    "2. **Interpreted language**: Python code is executed line by line, which makes debugging easier.\n"
    "3. **Cross-platform**: Python can run on various platforms such as Windows, Mac, Linux, etc.\n"
    "4. **Large standard library**: Python comes with a vast standard library, providing many pre-built functions.\n"
    "5. **Open-source**: Python is open-source and has a large community for support.\n"
    "6. **Dynamic typing**: Python allows dynamic typing, meaning variables can change types dynamically during execution."
)

# Question 2
doc.add_heading('2. Explain if…else statement with example.', level=2)
doc.add_paragraph(
    "The `if…else` statement in Python is used for conditional execution. It allows a program to perform one action if a condition is true, "
    "and a different action if the condition is false.\n\n"
    "Example:\n"
    "```python\n"
    "x = 10\n"
    "if x > 5:\n"
    "    print('x is greater than 5')\n"
    "else:\n"
    "    print('x is less than or equal to 5')\n"
    "```\n"
    "In this example, since `x` is 10, the output will be 'x is greater than 5'."
)

# Saving the document as PDF
doc.save('/mnt/data/Python_Question_Bank_with_Answers.docx')
